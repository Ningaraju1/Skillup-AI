import os
import re
import pypdf
import docx


def extract_text(file_path: str) -> str:
    """
    Extract text from a PDF, DOCX, or DOC file safely.
    Returns cleaned plain text.
    """
    if not file_path:
        return ""

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".doc":
        return extract_text_from_doc(file_path)
    else:
        return f"ERROR: Unsupported file extension '{ext}'."


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Safely extract text from PDF files using pypdf.
    """
    try:
        reader = pypdf.PdfReader(pdf_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"ERROR: Failed to extract PDF text - {str(e)}"


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract paragraphs and table contents from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(docx_path)
        full_text = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text).strip()
    except Exception as e:
        return f"ERROR: Failed to extract DOCX text - {str(e)}"


def extract_text_from_doc(doc_path: str) -> str:
    """
    Extract readable text streams from a legacy binary .doc file using binary inspection.
    This parses UTF-16-LE and printable ASCII blocks.
    """
    try:
        with open(doc_path, "rb") as f:
            content = f.read()

        # Extract UTF-16-LE streams (dominant format for Word text content)
        # Sequence of printable characters in UTF-16-LE (char + \x00)
        utf16_pattern = re.compile(rb"(?:[\x20-\x7E\r\n\t]\x00){4,}")
        utf16_matches = utf16_pattern.findall(content)
        utf16_text = ""
        if utf16_matches:
            decoded_parts = []
            for part in utf16_matches:
                try:
                    decoded_parts.append(part.decode("utf-16-le"))
                except Exception:
                    pass
            utf16_text = "\n".join(decoded_parts)

        # Extract ASCII streams as a fallback
        ascii_pattern = re.compile(rb"[\x20-\x7E\r\n\t]{4,}")
        ascii_matches = ascii_pattern.findall(content)
        ascii_text = ""
        if ascii_matches:
            decoded_parts = []
            for part in ascii_matches:
                try:
                    decoded = part.decode("ascii", errors="ignore")
                    # Exclude common OLE metadata headers to minimize noise
                    if any(
                        term in decoded
                        for term in [
                            "WordDocument",
                            "SummaryInformation",
                            "DocumentSummaryInformation",
                            "CompObj",
                            "ObjectPool",
                        ]
                    ):
                        continue
                    decoded_parts.append(decoded)
                except Exception:
                    pass
            ascii_text = "\n".join(decoded_parts)

        # Choose the stream that yielded the most content
        text = utf16_text if len(utf16_text) > len(ascii_text) * 0.5 else ascii_text

        # Clean up excessive whitespace and line breaks
        cleaned_text = re.sub(r"\n+", "\n", text)
        return cleaned_text.strip()
    except Exception as e:
        return f"ERROR: Failed to extract legacy DOC text - {str(e)}"